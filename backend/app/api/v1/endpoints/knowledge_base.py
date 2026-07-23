"""
Knowledge Base API Endpoints

Endpoints for managing knowledge bases, documents, and semantic search.
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime

from pydantic import BaseModel, Field

from app.core.dependencies import get_db, get_current_active_user
from app.models.user import User, OrganizationMember


async def get_current_org_id(
    user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
) -> uuid.UUID:
    """Resolve the user's organization id. User has no organization_id column;
    membership lives in organization_members. Falls back to the user id."""
    m = (await db.execute(
        select(OrganizationMember).where(OrganizationMember.user_id == user.id).limit(1)
    )).scalar_one_or_none()
    return m.organization_id if m else user.id
from app.models.knowledge_base import KnowledgeBase as KnowledgeBaseModel, Document as DocumentModel
from app.services.knowledge_base import RAGService
from app.core.config import settings

router = APIRouter()


# Pydantic Schemas
class KnowledgeBaseCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=255)
    description: Optional[str] = None
    chunk_size: int = Field(default=1000, ge=100, le=4000)
    chunk_overlap: int = Field(default=200, ge=0, le=1000)
    vector_store_type: str = Field(default="local")  # local, pinecone, qdrant
    vector_store_config: Optional[Dict[str, Any]] = None


class KnowledgeBaseResponse(BaseModel):
    id: uuid.UUID
    organization_id: uuid.UUID
    name: str
    description: Optional[str]
    embedding_model: str
    chunk_size: int
    chunk_overlap: int
    vector_store_type: str
    is_active: bool
    created_at: datetime
    updated_at: datetime
    document_count: int = 0

    class Config:
        from_attributes = True


class DocumentCreate(BaseModel):
    knowledge_base_id: uuid.UUID
    title: str = Field(..., min_length=1, max_length=500)
    content: str = Field(..., min_length=1)
    source_type: str = Field(default="text")  # text, file, url, api
    source_url: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class DocumentResponse(BaseModel):
    id: uuid.UUID
    knowledge_base_id: uuid.UUID
    title: str
    source_type: str
    source_url: Optional[str]
    file_type: Optional[str]
    file_size: Optional[int]
    processing_status: str
    processing_error: Optional[str]
    total_chunks: int
    total_tokens: Optional[int]
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class SearchRequest(BaseModel):
    query: str = Field(..., min_length=1)
    knowledge_base_id: uuid.UUID
    top_k: int = Field(default=5, ge=1, le=20)
    # Cosine scores from text-embedding-3-small for a genuine paraphrase match
    # land around 0.3-0.5, so a 0.7 floor silently discarded every real hit.
    min_similarity: float = Field(default=0.2, ge=0.0, le=1.0)
    filters: Optional[Dict[str, Any]] = None


class SearchResult(BaseModel):
    chunk_id: str
    content: str
    document_id: str
    document_title: str
    chunk_index: int
    similarity: float
    metadata: Dict[str, Any]


class SearchResponse(BaseModel):
    query: str
    results: List[SearchResult]
    total_results: int


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF. Image-only (scanned) PDFs yield nothing."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            status_code=400,
            detail="PDF support is not installed on the server (pypdf missing).",
        )

    import io

    try:
        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {e}")


def _extract_docx_text(content: bytes) -> str:
    """Extract text from a .docx, including table cells."""
    try:
        import docx
    except ImportError:
        raise HTTPException(
            status_code=400,
            detail="DOCX support is not installed on the server (python-docx missing).",
        )

    import io

    try:
        document = docx.Document(io.BytesIO(content))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read DOCX: {e}")


# Helper function to get RAG service
def get_rag_service(db: AsyncSession = Depends(get_db)) -> RAGService:
    """Get RAG service instance."""
    return RAGService(
        db=db,
        openai_api_key=settings.OPENAI_API_KEY,
        vector_store_config={
            'type': 'local',  # Can be configured via environment
            'config': {}
        }
    )


# Knowledge Base Endpoints

@router.post("/knowledge-bases", response_model=KnowledgeBaseResponse, status_code=201)
async def create_knowledge_base(
    kb_create: KnowledgeBaseCreate,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service)
):
    """
    Create a new knowledge base.

    Creates a knowledge base container for organizing documents.
    Automatically creates a vector store index.
    """
    try:
        kb = await rag_service.create_knowledge_base(
            organization_id=org_id,
            name=kb_create.name,
            description=kb_create.description,
            chunk_size=kb_create.chunk_size,
            chunk_overlap=kb_create.chunk_overlap,
            vector_store_type=kb_create.vector_store_type,
            vector_store_config=kb_create.vector_store_config or {}
        )

        return KnowledgeBaseResponse(
            **kb.__dict__,
            document_count=0
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/knowledge-bases", response_model=List[KnowledgeBaseResponse])
async def list_knowledge_bases(
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db)
):
    """
    List all knowledge bases for the organization.
    """
    from sqlalchemy import select, func
    from app.models.knowledge_base import Document

    result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.organization_id == org_id
        ).order_by(KnowledgeBaseModel.created_at.desc())
    )
    knowledge_bases = result.scalars().all()

    # Get document counts
    response_list = []
    for kb in knowledge_bases:
        count_result = await db.execute(
            select(func.count(Document.id)).where(Document.knowledge_base_id == kb.id)
        )
        doc_count = count_result.scalar() or 0

        response_list.append(
            KnowledgeBaseResponse(
                **kb.__dict__,
                document_count=doc_count
            )
        )

    return response_list


@router.get("/knowledge-bases/{kb_id}", response_model=KnowledgeBaseResponse)
async def get_knowledge_base(
    kb_id: uuid.UUID,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db)
):
    """
    Get a specific knowledge base by ID.
    """
    from sqlalchemy import select, func
    from app.models.knowledge_base import Document

    result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == kb_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Get document count
    count_result = await db.execute(
        select(func.count(Document.id)).where(Document.knowledge_base_id == kb.id)
    )
    doc_count = count_result.scalar() or 0

    return KnowledgeBaseResponse(
        **kb.__dict__,
        document_count=doc_count
    )


@router.delete("/knowledge-bases/{kb_id}", status_code=204)
async def delete_knowledge_base(
    kb_id: uuid.UUID,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db)
):
    """
    Delete a knowledge base and all its documents.
    """
    from sqlalchemy import select

    result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == kb_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Delete index from vector store
    # TODO: Implement vector store cleanup

    await db.delete(kb)
    await db.commit()

    return None


# Document Endpoints

@router.post("/documents", response_model=DocumentResponse, status_code=201)
async def create_document(
    doc_create: DocumentCreate,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service),
    db: AsyncSession = Depends(get_db)
):
    """
    Add a text document to a knowledge base.

    The document will be automatically chunked, embedded, and indexed.
    """
    # Verify knowledge base belongs to user's organization
    from sqlalchemy import select

    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == doc_create.knowledge_base_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    try:
        doc = await rag_service.add_document(
            knowledge_base_id=doc_create.knowledge_base_id,
            title=doc_create.title,
            content=doc_create.content,
            source_type=doc_create.source_type,
            source_url=doc_create.source_url,
            metadata=doc_create.metadata or {}
        )

        return DocumentResponse(**doc.__dict__)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/documents/upload", response_model=DocumentResponse, status_code=201)
async def upload_document(
    knowledge_base_id: uuid.UUID = Form(...),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service),
    db: AsyncSession = Depends(get_db)
):
    """
    Upload a file document to a knowledge base.

    Supported formats: txt, pdf, docx, md, json
    """
    # Verify knowledge base
    from sqlalchemy import select

    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == knowledge_base_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Read file content
    content = await file.read()

    # Extract text based on file type
    try:
        name = (file.filename or "").lower()

        if name.endswith(('.txt', '.md', '.json', '.csv')):
            text_content = content.decode('utf-8')
        elif name.endswith('.pdf'):
            text_content = _extract_pdf_text(content)
        elif name.endswith('.docx'):
            text_content = _extract_docx_text(content)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file format. Upload a .txt, .md, .json, .csv, .pdf or .docx file.",
            )

        if not text_content.strip():
            raise HTTPException(
                status_code=400,
                detail=(
                    "No text could be extracted from this file. Scanned/image-only "
                    "PDFs need OCR, which isn't supported — upload a text-based file."
                ),
            )

        doc = await rag_service.add_document(
            knowledge_base_id=knowledge_base_id,
            title=file.filename,
            content=text_content,
            source_type="file",
            file_type=file.content_type,
            file_size=len(content)
        )

        return DocumentResponse(**doc.__dict__)

    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File encoding error. Please upload UTF-8 encoded files.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/knowledge-bases/{kb_id}/documents", response_model=List[DocumentResponse])
async def list_documents(
    kb_id: uuid.UUID,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db)
):
    """
    List all documents in a knowledge base.
    """
    from sqlalchemy import select

    # Verify knowledge base
    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == kb_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Get documents
    result = await db.execute(
        select(DocumentModel).where(
            DocumentModel.knowledge_base_id == kb_id
        ).order_by(DocumentModel.created_at.desc())
    )
    documents = result.scalars().all()

    return [DocumentResponse(**doc.__dict__) for doc in documents]


@router.delete("/documents/{doc_id}", status_code=204)
async def delete_document(
    doc_id: uuid.UUID,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service),
    db: AsyncSession = Depends(get_db)
):
    """
    Delete a document and all its chunks.
    """
    from sqlalchemy import select

    # Verify document belongs to user's organization
    doc_result = await db.execute(
        select(DocumentModel).where(DocumentModel.id == doc_id)
    )
    doc = doc_result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Verify knowledge base ownership
    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == doc.knowledge_base_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=403, detail="Access denied")

    # Delete document
    success = await rag_service.delete_document(doc_id)

    if not success:
        raise HTTPException(status_code=500, detail="Failed to delete document")

    return None


# Search Endpoints

@router.post("/search", response_model=SearchResponse)
async def search_knowledge_base(
    search_req: SearchRequest,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service),
    db: AsyncSession = Depends(get_db)
):
    """
    Semantic search across a knowledge base.

    Uses embeddings to find relevant content based on semantic similarity.
    """
    from sqlalchemy import select

    # Verify knowledge base
    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == search_req.knowledge_base_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Search the embeddings stored in Postgres (DocumentChunk.embedding).
    # The RAGService vector-store path is process-local and starts empty on
    # every request, so it can never return a hit for previously-added
    # documents — the DB is the only durable home for these vectors.
    from app.services.knowledge_base.rag_service import search_knowledge_base_db

    hits = await search_knowledge_base_db(
        db=db,
        knowledge_base_id=str(search_req.knowledge_base_id),
        query=search_req.query,
        api_key=settings.OPENAI_API_KEY,
        top_k=search_req.top_k,
        min_similarity=search_req.min_similarity,
    )

    return SearchResponse(
        query=search_req.query,
        results=[
            SearchResult(
                chunk_id=str(h.get("chunk_id", "")),
                content=h.get("content", ""),
                document_id=str(h.get("document_id", "")),
                document_title=h.get("document_title") or "",
                chunk_index=h.get("chunk_index", 0),
                similarity=h.get("score", 0.0),
                metadata=h.get("metadata") or {},
            )
            for h in hits
        ],
        total_results=len(hits),
    )


@router.post("/search/context", response_model=Dict[str, str])
async def get_search_context(
    search_req: SearchRequest,
    max_tokens: int = 2000,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    rag_service: RAGService = Depends(get_rag_service),
    db: AsyncSession = Depends(get_db)
):
    """
    Get formatted context for LLM prompt injection.

    Returns a context string ready to be injected into prompts.
    """
    from sqlalchemy import select

    # Verify knowledge base
    kb_result = await db.execute(
        select(KnowledgeBaseModel).where(
            KnowledgeBaseModel.id == search_req.knowledge_base_id,
            KnowledgeBaseModel.organization_id == org_id
        )
    )
    kb = kb_result.scalar_one_or_none()

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Get context
    context = await rag_service.get_context_for_prompt(
        knowledge_base_id=search_req.knowledge_base_id,
        query=search_req.query,
        max_chunks=search_req.top_k,
        max_tokens=max_tokens
    )

    return {
        "context": context,
        "query": search_req.query
    }


# ============================================================================
# Agent <-> Knowledge Base attachment
# ============================================================================


class AgentKnowledgeBaseLink(BaseModel):
    """A knowledge base attached to an agent."""
    knowledge_base_id: uuid.UUID
    name: Optional[str] = None
    max_results: int = 5
    min_similarity: float = 0.2
    auto_inject: bool = True
    is_active: bool = True


class AgentKnowledgeBaseUpdate(BaseModel):
    """Replace the set of knowledge bases attached to an agent."""
    knowledge_base_ids: List[uuid.UUID] = Field(default_factory=list)
    max_results: int = Field(default=5, ge=1, le=20)
    min_similarity: float = Field(default=0.2, ge=0.0, le=1.0)
    auto_inject: bool = True


@router.get("/agents/{agent_id}/knowledge-bases", response_model=List[AgentKnowledgeBaseLink])
async def get_agent_knowledge_bases(
    agent_id: uuid.UUID,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db),
):
    """List the knowledge bases attached to an agent."""
    from app.models.knowledge_base import AgentKnowledgeBase

    rows = (
        await db.execute(
            select(AgentKnowledgeBase, KnowledgeBaseModel)
            .join(
                KnowledgeBaseModel,
                KnowledgeBaseModel.id == AgentKnowledgeBase.knowledge_base_id,
            )
            .where(
                AgentKnowledgeBase.agent_id == agent_id,
                KnowledgeBaseModel.organization_id == org_id,
            )
            .order_by(AgentKnowledgeBase.priority.desc())
        )
    ).all()

    return [
        AgentKnowledgeBaseLink(
            knowledge_base_id=link.knowledge_base_id,
            name=kb.name,
            max_results=link.max_results,
            min_similarity=link.min_similarity,
            auto_inject=link.auto_inject,
            is_active=link.is_active,
        )
        for link, kb in rows
    ]


@router.put("/agents/{agent_id}/knowledge-bases", response_model=List[AgentKnowledgeBaseLink])
async def set_agent_knowledge_bases(
    agent_id: uuid.UUID,
    payload: AgentKnowledgeBaseUpdate,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db),
):
    """
    Replace the set of knowledge bases attached to an agent.

    Sending an empty list detaches everything. Only knowledge bases owned by the
    caller's organization can be attached.
    """
    from app.models.agent import Agent
    from app.models.knowledge_base import AgentKnowledgeBase

    agent = (
        await db.execute(select(Agent).where(Agent.id == agent_id))
    ).scalar_one_or_none()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    requested = list(dict.fromkeys(payload.knowledge_base_ids))  # de-dupe, keep order

    if requested:
        owned = (
            await db.execute(
                select(KnowledgeBaseModel.id).where(
                    KnowledgeBaseModel.id.in_(requested),
                    KnowledgeBaseModel.organization_id == org_id,
                )
            )
        ).scalars().all()
        missing = set(requested) - set(owned)
        if missing:
            raise HTTPException(
                status_code=404,
                detail=f"Knowledge base(s) not found: {', '.join(str(m) for m in missing)}",
            )

    # Replace wholesale — simplest correct semantics for a checkbox-style UI.
    existing = (
        await db.execute(
            select(AgentKnowledgeBase).where(AgentKnowledgeBase.agent_id == agent_id)
        )
    ).scalars().all()
    for row in existing:
        await db.delete(row)

    for i, kb_id in enumerate(requested):
        db.add(
            AgentKnowledgeBase(
                agent_id=agent_id,
                knowledge_base_id=kb_id,
                priority=len(requested) - i,  # preserve the given order
                max_results=payload.max_results,
                min_similarity=payload.min_similarity,
                auto_inject=payload.auto_inject,
                is_active=True,
            )
        )

    await db.commit()

    return await get_agent_knowledge_bases(
        agent_id=agent_id, current_user=current_user, org_id=org_id, db=db
    )


# ============================================================================
# Ask (retrieve + synthesize) — what the caller actually hears
# ============================================================================


class AskRequest(BaseModel):
    knowledge_base_id: uuid.UUID
    query: str = Field(..., min_length=1)
    top_k: int = Field(default=5, ge=1, le=20)
    min_similarity: float = Field(default=0.2, ge=0.0, le=1.0)


class AskSource(BaseModel):
    document_title: str
    content: str
    similarity: float


class AskResponse(BaseModel):
    query: str
    answer: str
    sources: List[AskSource]


@router.post("/ask", response_model=AskResponse)
async def ask_knowledge_base(
    req: AskRequest,
    current_user: User = Depends(get_current_active_user),
    org_id: uuid.UUID = Depends(get_current_org_id),
    db: AsyncSession = Depends(get_db),
):
    """
    Answer a question from the knowledge base.

    Retrieves the relevant passages, then has the LLM write a direct answer from
    them — the same retrieve-then-generate path an agent uses on a live call, so
    this previews what a caller would actually hear.
    """
    from app.services.knowledge_base.rag_service import search_knowledge_base_db

    kb = (
        await db.execute(
            select(KnowledgeBaseModel).where(
                KnowledgeBaseModel.id == req.knowledge_base_id,
                KnowledgeBaseModel.organization_id == org_id,
            )
        )
    ).scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    hits = await search_knowledge_base_db(
        db=db,
        knowledge_base_id=str(req.knowledge_base_id),
        query=req.query,
        api_key=settings.OPENAI_API_KEY,
        top_k=req.top_k,
        min_similarity=req.min_similarity,
    )

    sources = [
        AskSource(
            document_title=h.get("document_title") or "document",
            content=h.get("content", ""),
            similarity=h.get("score", 0.0),
        )
        for h in hits
    ]

    if not hits:
        return AskResponse(
            query=req.query,
            answer=(
                "I don't have anything about that in this knowledge base. "
                "Try wording the question closer to your documents, or add a "
                "document that covers it."
            ),
            sources=[],
        )

    context = "\n\n---\n\n".join(
        f"[{h.get('document_title') or 'document'}]\n{h.get('content', '')}" for h in hits
    )

    try:
        from app.services.voice.llm_service import get_llm_service, ChatMessage

        llm = get_llm_service()
        completion = await llm.chat(
            messages=[
                ChatMessage(
                    role="system",
                    content=(
                        "You answer questions using ONLY the provided context from the "
                        "company knowledge base. Give a direct, natural answer in a few "
                        "sentences — the way you would say it out loud on a phone call. "
                        "Do not use markdown, bullet points, or headings. If the context "
                        "does not contain the answer, say you don't have that information."
                    ),
                ),
                ChatMessage(
                    role="user",
                    content=f"Context:\n{context}\n\nQuestion: {req.query}",
                ),
            ],
            model="gpt-4o-mini",
            temperature=0.2,
            max_tokens=300,
        )
        answer = (getattr(completion, "content", None) or "").strip()
    except Exception as e:
        # Retrieval still succeeded — surface the passages rather than 500ing.
        import logging

        logging.getLogger(__name__).error(f"Answer synthesis failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=502,
            detail=f"Found relevant passages but could not generate an answer: {e}",
        )

    if not answer:
        answer = "I couldn't produce an answer from the retrieved content."

    return AskResponse(query=req.query, answer=answer, sources=sources)
